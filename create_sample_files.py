import os
import docx

def create_sample_docx_resume():
    """
    Generates a sample resume DOCX file in the workspace directory.
    """
    doc = docx.Document()
    
    # Title / Header
    doc.add_heading("Alex Mercer", level=0)
    p = doc.add_paragraph("Email: alex.mercer@gmail.com | Phone: (555) 019-2834 | GitHub: github.com/alex-mercer")
    
    # Summary
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(
        "Highly motivated software engineer with 2+ years of professional experience designing, building, "
        "and scaling RESTful web APIs and frontend web applications. Expert in Python, SQL databases, "
        "and frontend integrations."
    )
    
    # Skills
    doc.add_heading("Core Technical Skills", level=1)
    doc.add_paragraph(
        "Languages: Python, SQL, JavaScript, HTML, CSS\n"
        "Frameworks & Libraries: Flask, Django, PostgreSQL, SQLite, React, Tailwind CSS\n"
        "Tools & Workflows: Docker, Git, RESTful APIs, Agile/Scrum"
    )
    
    # Experience
    doc.add_heading("Professional Experience", level=1)
    
    doc.add_heading("Software Engineer | Innovate Solutions", level=2)
    doc.add_paragraph("January 2024 - Present | Austin, TX")
    doc.add_paragraph("Lead full stack application development for core customer portal services.")
    doc.add_paragraph("Key Achievements:")
    doc.add_paragraph("• Developed and scaling serverless REST API endpoints using Python Flask, decreasing response times by 35%.")
    doc.add_paragraph("• Managed relational database migrations to PostgreSQL, optimizing lookup query indices.")
    
    doc.add_heading("Junior Developer | Code Labs", level=2)
    doc.add_paragraph("June 2022 - December 2023 | Dallas, TX")
    doc.add_paragraph("Assisted in backend codebase maintenance and React frontend user interactions.")
    doc.add_paragraph("Key Achievements:")
    doc.add_paragraph("• Resolved hook dependencies in React modules, preventing redundant component re-renders.")
    doc.add_paragraph("• Wrote comprehensive unit tests using PyTest to achieve 85% coverage across endpoints.")
    
    # Education
    doc.add_heading("Education", level=1)
    doc.add_heading("Bachelor of Science in Computer Science", level=2)
    doc.add_paragraph("Tech University | Graduated May 2022")
    
    filename = "sample_resume.docx"
    doc.save(filename)
    print(f"Sample DOCX resume successfully created: {filename}")

def create_sample_job_description():
    """
    Generates a sample target job description text file.
    """
    jd_text = """Job Title: Senior Backend Developer
Company: Apex Cloud Analytics
Location: Remote (US)
Employment Type: Full-Time

Job Summary:
Apex Cloud Analytics is looking for a Senior Backend Developer to join our core team. You will be responsible for designing and scaling asynchronous microservices, developing database schemas, and maintaining high performance REST APIs that handle millions of transactions weekly.

Core Responsibilities:
- Architect, build, and optimize backend microservices using Python (Flask or FastAPI).
- Integrate Redis caching layers to reduce database latency.
- Containerize application workloads using Docker and deploy into Kubernetes clusters.
- Collaborate with frontend engineers to construct clean, reusable API contracts.
- Design database architectures and manage PostgreSQL migrations.

Required Skills & Experience:
- 5+ years of software development experience (primarily with Python).
- Deep experience in designing scalable RESTful APIs.
- Production experience containerizing applications using Docker.
- Advanced proficiency in SQL and PostgreSQL query tuning.
- Knowledge of Git and CI/CD automation pipelines.
- Experience with TypeScript, React, and frontend setups is a plus.
"""
    filename = "sample_job_description.txt"
    with open(filename, "w", encoding="utf-8") as f:
        f.write(jd_text.strip())
    print(f"Sample Job Description successfully created: {filename}")

if __name__ == "__main__":
    create_sample_docx_resume()
    create_sample_job_description()
