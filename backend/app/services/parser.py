import io
import pdfplumber
import docx

def parse_pdf(file_bytes: bytes) -> str:
    """
    Extracts text from PDF bytes using pdfplumber.
    """
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        raise ValueError(f"Error parsing PDF file: {str(e)}")
    return text

def parse_docx(file_bytes: bytes) -> str:
    """
    Extracts text from DOCX bytes using python-docx.
    """
    text = ""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + " "
                text += "\n"
    except Exception as e:
        raise ValueError(f"Error parsing DOCX file: {str(e)}")
    return text

def parse_document(filename: str, file_bytes: bytes) -> str:
    """
    Auto-detects format from filename extension and extracts plain text content.
    """
    ext = filename.lower().split('.')[-1]
    if ext == 'pdf':
        return parse_pdf(file_bytes)
    elif ext in ['docx', 'doc']:
        return parse_docx(file_bytes)
    else:
        raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")
Base = None
